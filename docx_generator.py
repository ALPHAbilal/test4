"""
DOCX Generator Module

This module provides functionality to convert markdown content to DOCX format.
"""

import os
import uuid
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import re

# Setup logging
logger = logging.getLogger(__name__)

# Constants
DOCX_OUTPUT_DIR = os.getenv('DOCX_OUTPUT_DIR', 'docx_output')

# Ensure the path is absolute
if not os.path.isabs(DOCX_OUTPUT_DIR):
    DOCX_OUTPUT_DIR = os.path.abspath(DOCX_OUTPUT_DIR)

# Ensure the output directory exists
os.makedirs(DOCX_OUTPUT_DIR, exist_ok=True)
logger.info(f"DOCX output directory set to: {DOCX_OUTPUT_DIR}")

def markdown_to_docx(markdown_content, template_name):
    """
    Converts markdown content to a professionally formatted DOCX file.

    Args:
        markdown_content: The populated markdown content
        template_name: Name of the original template for reference

    Returns:
        Tuple of (file_path, file_name) for the generated DOCX
    """
    logger.info(f"Converting markdown to DOCX for template: {template_name}")

    # Create a unique filename based on the template name
    base_name = os.path.splitext(template_name)[0]
    file_name = f"{base_name}_{uuid.uuid4().hex[:8]}.docx"
    file_path = os.path.join(DOCX_OUTPUT_DIR, file_name)

    # Convert markdown to HTML
    html_content = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Document
    doc = Document()

    # Set document properties
    doc.core_properties.title = base_name
    doc.core_properties.author = "Document Generator"

    # Add a title
    title = doc.add_heading(base_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process HTML elements
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table']):
        if element.name.startswith('h'):
            # Handle headings
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(), level=level)

        elif element.name == 'p':
            # Handle paragraphs
            paragraph = doc.add_paragraph()

            # Process inline elements
            for content in element.contents:
                if content.name == 'strong' or content.name == 'b':
                    paragraph.add_run(content.get_text()).bold = True
                elif content.name == 'em' or content.name == 'i':
                    paragraph.add_run(content.get_text()).italic = True
                elif content.name == 'a':
                    paragraph.add_run(content.get_text()).underline = True
                else:
                    paragraph.add_run(str(content))

        elif element.name == 'ul' or element.name == 'ol':
            # Handle lists
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph(li.get_text())
                paragraph.style = 'List Bullet' if element.name == 'ul' else 'List Number'

        elif element.name == 'table':
            # Handle tables
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()

    # Save the document
    doc.save(file_path)
    logger.info(f"DOCX file saved to: {file_path}")

    return file_path, file_name